import os
import logging
from io import BytesIO
import streamlit as st
import PyPDF2
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from docx import Document
from datetime import datetime
import re

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger()

# Load environment variables
load_dotenv()

# Funzione per configurare l'API OpenAI e inizializzare il modello
def openai_m():
    api_choice = st.sidebar.selectbox("Scegli la chiave API da usare", ["Usa chiave di sistema", "Inserisci la tua chiave API"], index=1)
    
    if api_choice == "Inserisci la tua chiave API":
        openai_api_key = st.sidebar.text_input("Inserisci la tua chiave API OpenAI", st.session_state.get("user_api_key", ""), type="password")
        st.session_state.user_api_key = openai_api_key
    else:
        openai_api_key = os.getenv("OPENAI_API_KEY")
    
    if not openai_api_key:
        st.error("Errore: La chiave API non è stata inserita o non è configurata correttamente.")
        return None, None, None

    model_choice = st.sidebar.selectbox("Seleziona il modello LLM", ["gpt-4o", "gpt-4o-mini"], index=1)
    temperature = st.sidebar.slider("Imposta la temperatura del modello", min_value=0.0, max_value=1.0, value=0.7, step=0.1)

    return openai_api_key, model_choice, temperature

# Funzioni per l'estrazione del testo
def extract_text_from_pdf(reader):
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n\n"
    return text

def extract_text_from_doc(doc_file):
    logger.info("Extracting text from DOC file.")
    doc = Document(doc_file)
    text = "\n\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_txt(txt_file):
    logger.info("Extracting text from TXT file.")
    text = txt_file.read().decode("utf-8")
    return text

# Pulizia della formattazione Markdown
def clean_markdown_formatting(text):
    """
    Rimuove la formattazione Markdown dal testo.
    """
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Rimuove grassetto
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Rimuove corsivo
    text = re.sub(r'\#\#(.*?)\n', r'\1\n', text)  # Rimuove titoli di secondo livello
    text = re.sub(r'\#(.*?)\n', r'\1\n', text)    # Rimuove titoli di primo livello
    text = re.sub(r'[-*]\s', '', text)            # Rimuove elenchi puntati
    return text

# Caricamento e pulizia del file
def upload_and_extract_text():
    uploaded_file = st.file_uploader("Carica un file PDF, DOCX o TXT", type=["pdf", "docx", "txt"])

    if uploaded_file is not None:
        file_type = uploaded_file.name.split('.')[-1].lower()
        logger.info(f"Uploaded file: {uploaded_file.name}")

        if file_type == "pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            text = extract_text_from_pdf(reader)
        elif file_type == "docx":
            text = extract_text_from_doc(uploaded_file)
        elif file_type == "txt":
            text = extract_text_from_txt(uploaded_file)
        else:
            st.error("Formato file non supportato.")
            return None

        return clean_markdown_formatting(text)
    else:
        return None

# Divisione del testo in parti
def split_text_by_delimiter(text, delimiter="+++++"):
    return text.split(delimiter)

# Funzione per generare domande a scelta multipla usando il modello OpenAI
def generate_questions(text_segment, llm):
    prompt_template = """
    Genera una domanda a scelta multipla basata sul seguente testo:

    {text}

    Riferisciti al testo come unità didattica. La domanda deve avere 4 opzioni di risposta, la cui lunghezza deve essere molto simile, una delle quali è corretta. Ogni opzione deve essere plausibile e coerente con il contenuto dell'Unità didattica.

    L'obiettivo principale è far riflettere gli studenti su concetti chiave dell'Unità didattica, stimolando la loro capacità di analizzare e applicare tali concetti in situazioni reali.

    Per ogni risposta, sia giusta che errata, fornisci una spiegazione dettagliata seguendo queste indicazioni:
    
    - Le opzioni di risposta (inclusi i distrattori) devono essere formulate in modo plausibile, evitando banalità. I distrattori devono essere basati su potenziali errori concettuali che uno studente avanzato potrebbe commettere, ma che portino a riflettere sui concetti fondamentali del testo.
    - Non utilizzare frasi come "il testo non menziona", ma fornisci una spiegazione plausibile legata al contesto dell'Unità didattica.
    - La spiegazione della risposta corretta deve essere esauriente, chiara e direttamente collegata al concetto chiave trattato nell'Unità didattica.
    - Mantieni un linguaggio formale e preciso, ma chiaro, adeguato per studenti di livello avanzato, come quelli di un master. Evita tecnicismi eccessivi, a meno che non siano essenziali per la comprensione del concetto trattato.
    """
    
    prompt = ChatPromptTemplate.from_template(template=prompt_template)
    prompt_input = prompt.format(text=text_segment)

    # Chiamata al modello LLM per generare la domanda
    response = llm(prompt_input)

    # Debug: Esamina la risposta dal modello
    st.write("Risposta dal modello:", response)

    # Verifica se la risposta è un oggetto AIMessage e ritorna il contenuto senza altre elaborazioni
    if hasattr(response, 'content'):
        st.write("La risposta è un oggetto AIMessage. Riportiamo il contenuto così com'è.")
        return clean_markdown_formatting(response.content)  # Ripuliamo il testo dalla formattazione Markdown
    else:
        st.error(f"Formato della risposta non previsto: {type(response)}")
        return None

# Creazione di un file DOCX con il contenuto
def write_questions_to_docx(questions, output_filename):
    doc = Document()
    
    for i, question_data in enumerate(questions):
        if question_data is None:
            continue
        
        # Scriviamo il contenuto intero di ogni segmento senza alcuna modifica
        doc.add_paragraph(f"{i+1}. -------------------------")
        doc.add_paragraph(question_data)
        doc.add_paragraph("\n")
    
    # Salvataggio del file DOCX
    with open(output_filename, "wb") as f:
        doc.save(f)

# Funzione principale per inizializzare il modello, caricare il file, generare contenuti e creare il DOCX
def generate_questions_from_text():
    st.title("Generatore di Domande da Testo")
    
    # Inizializzazione del modello LLM
    openai_api_key, model_choice, temperature = openai_m()
    
    if not openai_api_key:
        st.error("Errore: la chiave API non è configurata correttamente.")
        return
    
    # Caricamento e estrazione del testo solo dopo l'inizializzazione del modello
    text = upload_and_extract_text()
    
    if text is not None:
        # Dividere il testo in parti utilizzando "+++++"
        text_segments = split_text_by_delimiter(text)

        # Configurazione del modello LLM
        llm = ChatOpenAI(
            temperature=temperature, 
            api_key=openai_api_key, 
            model_name=model_choice
        )

      # Genera il contenuto per ciascuna parte del testo usando un ciclo for
        questions = []
        for i, segment in enumerate(text_segments):
            st.write(f"Generando contenuto per il segmento {i+1}...")
            question_data = generate_questions(segment, llm)
            if question_data is not None:
                questions.append(question_data)

        # Creazione del nome del file DOCX di output con il nome del file di origine e data/ora
        # Ottieni la data e l'ora attuali per il timestamp
        day = datetime.now().strftime('%d')  # Giorno corrente
        hour = datetime.now().strftime('%H')  # Ora corrente

        # Genera il nome del file con il formato "domande_multiple_del_XX_ora_XX.docx"
        output_filename = f"domande_multiple_del_{day}_ora_{hour}.docx"

        # Salva il contenuto generato in un file DOCX
        write_questions_to_docx(questions, output_filename)

        # Mostra il link per il download del file DOCX
        st.success(f"File generato: {output_filename}")
        with open(output_filename, "rb") as f:
            st.download_button("Scarica il file DOCX generato", data=f, file_name=output_filename)


if __name__ == "__main__":
    generate_questions_from_text()
