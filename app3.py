import os
import logging
import streamlit as st
import PyPDF2
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from docx import Document
from datetime import datetime
import re

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger()

# Load environment variables
load_dotenv()

# Funzione per configurare l'API OpenAI e inizializzare il modello
def openai_m():
    api_choice = st.sidebar.selectbox("Scegli la chiave API da usare", ["Usa chiave di sistema", "Inserisci la tua chiave API"], index=1)
    
    if api_choice == "Inserisci la tua chiave API":
        openai_api_key = st.sidebar.text_input("Inserisci la tua chiave API OpenAI", st.session_state.get("user_api_key", ""), type="password")
        st.session_state.user_api_key = openai_api_key
    else:
        openai_api_key = os.getenv("OPENAI_API_KEY")
    
    if not openai_api_key:
        st.error("Errore: La chiave API non è stata inserita o non è configurata correttamente.")
        return None, None, None

    model_choice = st.sidebar.selectbox("Seleziona il modello LLM", ["gpt-4o", "gpt-4o-mini"], index=1)
    temperature = st.sidebar.slider("Imposta la temperatura del modello", min_value=0.0, max_value=1.0, value=0.7, step=0.1)

    return openai_api_key, model_choice, temperature

# Funzioni per l'estrazione del testo
def extract_text_from_pdf(reader):
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n\n"
    return text

def extract_text_from_doc(doc_file):
    logger.info("Extracting text from DOC file.")
    doc = Document(doc_file)
    text = "\n\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_txt(txt_file):
    logger.info("Extracting text from TXT file.")
    text = txt_file.read().decode("utf-8")
    return text

# Pulizia della formattazione Markdown
def clean_markdown_formatting(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Rimuove grassetto
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Rimuove corsivo
    text = re.sub(r'\#\#(.*?)\n', r'\1\n', text)  # Rimuove titoli di secondo livello
    text = re.sub(r'\#(.*?)\n', r'\1\n', text)    # Rimuove titoli di primo livello
    text = re.sub(r'[-*]\s', '', text)            # Rimuove elenchi puntati
    return text

# Caricamento e pulizia del file
def upload_and_extract_text():
    uploaded_file = st.file_uploader("Carica un file PDF, DOCX o TXT", type=["pdf", "docx", "txt"])

    if uploaded_file is not None:
        file_type = uploaded_file.name.split('.')[-1].lower()
        logger.info(f"Uploaded file: {uploaded_file.name}")

        if file_type == "pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            text = extract_text_from_pdf(reader)
        elif file_type == "docx":
            text = extract_text_from_doc(uploaded_file)
        elif file_type == "txt":
            text = extract_text_from_txt(uploaded_file)
        else:
            st.error("Formato file non supportato.")
            return None

        return clean_markdown_formatting(text)
    else:
        return None

# Divisione del testo in chunk
def split_text_into_chunks(text, num_chunks):
    chunk_size = len(text) // num_chunks
    chunks = []
    start = 0
    
    for i in range(num_chunks):
        end = start + chunk_size
        if i == num_chunks - 1:
            end = len(text)  # Assicuriamoci che l'ultimo chunk contenga il resto del testo
        chunk = text[start:end].strip()  # Rimuove spazi extra prima di aggiungere il chunk
        chunks.append(chunk)
        start = end
        logger.info(f"Chunk {i+1}: {chunk[:100]}...")  # Mostra solo i primi 100 caratteri per log
    return chunks

# Funzione per generare domande a scelta multipla usando il modello OpenAI
def generate_questions(text_segment, llm):
    prompt_template = """
    Genera una domanda a scelta multipla basata sul seguente testo: {text}. Segui le linee guida qui sotto per strutturare la domanda:

    - Stimolo chiaro: La domanda deve essere formulata in modo chiaro e preciso, testando un concetto chiave o una competenza specifica dell'unità didattica.
    - Livello di difficoltà: La domanda deve essere di livello [base/intermedio/avanzato], mirata a verificare [concetti di comprensione generale/applicazione pratica/sintesi critica].
    - Quattro opzioni di risposta: Fornisci quattro risposte alternative (a, b, c, d), una delle quali corretta. Tutte le risposte devono essere formulate in modo uniforme in termini di lunghezza, struttura grammaticale e complessità per evitare indizi involontari.
    - Distrattori plausibili: I tre distrattori devono rappresentare diversi tipi di errori concettuali, come fraintendimenti comuni, generalizzazioni eccessive o interpretazioni erronee. Ogni distrattore deve essere plausibile, ma chiaramente errato rispetto alla risposta corretta.
    - Concetti chiave: La domanda deve essere centrata su un concetto chiave ben definito all'interno del testo fornito, evitando ambiguità.
    - Tipo di domanda: Indica se la domanda verifica la comprensione teorica o la capacità di applicare il concetto in un contesto pratico.
    - Analisi della risposta corretta: Fornisci una spiegazione dettagliata e chiara che illustri perché la risposta corretta è giusta, collegandola direttamente al concetto chiave dell'unità didattica.
    - Analisi dei distrattori: Spiega perché ciascun distrattore è sbagliato o parzialmente corretto. Usa frasi come "è parzialmente corretto, ma..." o "sebbene sembri plausibile, manca di...". Spiega il tipo di errore concettuale commesso.

    Struttura della domanda:

    - Domanda: inserisci qui la domanda.
    - Risposte (una per riga):
        a) Risposta 1
        b) Risposta 2
        c) Risposta 3
        d) Risposta 4
    - Analisi della risposta corretta: inserisci qui l'analisi della risposta corretta.
    - Analisi dei distrattori: inserisci qui l'analisi dettagliata dei distrattori.

    Assicurati che tutte le risposte siano ben formulate e che i distrattori siano coerenti ma concettualmente errati.
    """
    
    prompt = ChatPromptTemplate.from_template(template=prompt_template)
    prompt_input = prompt.format(text=text_segment)

    response = llm(prompt_input)

    if hasattr(response, 'content'):
        return clean_markdown_formatting(response.content)
    else:
        st.error(f"Formato della risposta non previsto: {type(response)}")
        return None

# Creazione di un file DOCX con il contenuto
def write_questions_to_docx(questions, output_filename):
    doc = Document()
    for i, question_data in enumerate(questions):
        if question_data is not None:
            doc.add_paragraph(f"{i+1}. ")
            doc.add_paragraph(question_data)
            doc.add_paragraph("\n")
    doc.save(output_filename)

# Funzione principale per caricare il file e generare le domande
def generate_questions_from_text():
    st.title("Generatore di Domande da Testo")
    
    # Inizializzazione del modello LLM
    openai_api_key, model_choice, temperature = openai_m()
    
    if not openai_api_key:
        st.error("Errore: la chiave API non è configurata correttamente.")
        return
    
    text = upload_and_extract_text()
    
    if text is not None:
        # Numero di chunk deciso dall'utente
        num_chunks = st.number_input("In quanti pezzi vuoi dividere il testo?", min_value=1, max_value=50, value=5)
        
        if st.button("Avvia la suddivisione e la generazione delle domande"):
            # Dividi il testo in chunk
            chunks = split_text_into_chunks(text, num_chunks)
            st.success(f"Testo diviso in {len(chunks)} blocchi.")

            # Configurazione del modello LLM
            llm = ChatOpenAI(
                temperature=temperature, 
                api_key=openai_api_key, 
                model_name=model_choice
            )

            questions = []
            for i, segment in enumerate(chunks):
                st.write(f"Generando contenuto per il chunk {i+1}...")
                question_data = generate_questions(segment, llm)
                if question_data is not None:
                    questions.append(question_data)
                    
                    # Mostra a video le domande generate per il chunk
                    st.write(f"Domande per il chunk {i+1}:")
                    st.write(question_data)

            # Creazione del nome del file DOCX di output con la data e l'ora
            day = datetime.now().strftime('%d')
            hour = datetime.now().strftime('%H')
            output_filename = f"domande_multiple_del_{day}_ora_{hour}.docx"
            
            # Salva il contenuto generato in un file DOCX
            write_questions_to_docx(questions, output_filename)

            # Mostra il link per il download del file DOCX
            st.success(f"File generato: {output_filename}")
            with open(output_filename, "rb") as f:
                st.download_button("Scarica il file DOCX generato", data=f, file_name=output_filename)

if __name__ == "__main__":
    generate_questions_from_text()
